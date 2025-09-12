from docx import Document
from typing import List, Dict, Any
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from .models import (MultiLocationSchedule, WeeklySchedule, ScheduleEntry, 
                     DayOfWeek, TimeSlot, DetailCategory, TableCell, 
                     MultiLevelSchedule, SpecialityLevelSchedule, 
                     MultiStaffSchedule, StaffSchedule)
from enum import Enum
from datetime import datetime
import os


# Level and Speciality Mappings
LEVEL_MAPPING = {
    "100": "الأول",
    "200": "الثاني", 
    "300": "الثالث",
    "400": "الرابع"
}

# Base speciality mapping (for most cases)
SPECIALITY_MAPPING = {
    "pow": "القوي والآلات الكهربية",
    "comm": "الاتصالات",
    "comp": "الحاسبات"
}

# Level-dependent speciality mapping for special cases
LEVEL_SPECIALITY_MAPPING = {
    "comm": {
        "100": "الاتصالات والحاسبات",
        "200": "الاتصالات والحاسبات"
    }
}

# Header constants (moved to level config)

# Footer constants
FOOTER_GENERATION_INFO = "{date}"

# Base configuration for all levels
BASE_LEVEL_CONFIG = {
    "header": {
        "university_name": "جامعة الفيوم",
        "faculty_name": "كلية الهندسة",
        "department_prefix": "قسم الهندسة الكهربية",
        "academic_year": "العام الجامعي 2025 - 2026",
        "semester": "الفصل الدراسي الأول",
        "division_prefix": "شعبة",
        "level_prefix": "الفرقة",
        "schedule_template": "شئون التعليم والطلاب الجداول الدراسية"
    },
    "footer": {
        "dean_title": "عميد الكلية",
        "dean_name": "ا.د. رانيا احمد عبدالعظيم",
        "vice_dean_title": "وكيل الكلية لشئون التعليم والطلاب",
        "vice_dean_name": "ا.د. احمد سرج فريد",
        "head_department_title": "رئيس قسم الهندسة الكهربية",
        "head_department_name": "ا.د. عمرو رفعت",
    }
}

# Level-specific overrides
LEVEL_OVERRIDES = {
    "100": {
        "header": {
            "level_prefix": "المستوي"
        },
        "footer": {
            "program_manager_title": "منسق البرنامج",
        }
    },
    "200": {
        "header": {
            "level_prefix": "المستوي"
        },
        "footer": {
            "program_manager_title": "منسق البرنامج",
        }
    }
}

# Specialty-specific overrides
SPECIALTY_OVERRIDES = {
    "comm": {
        "footer": {
            "program_manager_name": "أ.د محمد حمدي"
        }
    },
    "pow": {
        "footer": {
            "program_manager_name": "أ.د خالد حسني"
        }
    }
}

# Function to get mapped value
def _get_mapped_level(level: str) -> str:
    """Get the Arabic representation of the level"""
    return LEVEL_MAPPING.get(level, level)

def _get_mapped_speciality(speciality: str, level: str = None) -> str:
    """Get the Arabic representation of the speciality"""
    # Check if this speciality has level-dependent mapping
    if speciality in LEVEL_SPECIALITY_MAPPING and level:
        level_mapping = LEVEL_SPECIALITY_MAPPING[speciality]
        if level in level_mapping:
            return level_mapping[level]
    
    # Fall back to base mapping
    return SPECIALITY_MAPPING.get(speciality, speciality)

# Function to merge base config with level-specific overrides
def _get_level_config(level: str, specialty: str = None) -> dict:
    """Get the configuration for a specific level and specialty by merging base config with overrides"""
    import copy
    
    # Start with a deep copy of the base configuration
    config = copy.deepcopy(BASE_LEVEL_CONFIG)
    
    # Apply level-specific overrides if they exist
    if level in LEVEL_OVERRIDES:
        level_override = LEVEL_OVERRIDES[level]
        
        # Merge header overrides
        if "header" in level_override:
            config["header"].update(level_override["header"])
        
        # Merge footer overrides
        if "footer" in level_override:
            config["footer"].update(level_override["footer"])
    
    # Apply specialty-specific overrides if they exist
    if specialty and specialty in SPECIALTY_OVERRIDES:
        specialty_override = SPECIALTY_OVERRIDES[specialty]
        
        # Merge header overrides
        if "header" in specialty_override:
            config["header"].update(specialty_override["header"])
        
        # Merge footer overrides
        if "footer" in specialty_override:
            config["footer"].update(specialty_override["footer"])
    
    return config

# Level-specific header and footer configurations (computed from base + overrides)
# Note: This is now deprecated in favor of calling _get_level_config directly with level and specialty
LEVEL_CONFIGS = {
    level: _get_level_config(level) for level in ["100", "200", "300", "400"]
}


class ColumnType(Enum):
    """Enum for different column types"""
    DAYS = 0
    CATEGORIES = 1
    TIME_SLOT_1 = 2
    TIME_SLOT_1_HALF = 3  # Half slot column for slot 1
    SEPARATOR_1 = 4
    TIME_SLOT_2 = 5
    TIME_SLOT_2_HALF = 6  # Half slot column for slot 2
    SEPARATOR_2 = 7
    TIME_SLOT_3 = 8
    TIME_SLOT_3_HALF = 9  # Half slot column for slot 3
    SEPARATOR_3 = 10
    TIME_SLOT_4 = 11
    TIME_SLOT_4_HALF = 12  # Half slot column for slot 4


class RowType(Enum):
    """Enum for different row types"""
    HEADER = 0
    DAY_START = 1  # First row of each day (course_name)
    DAY_MIDDLE = 2  # Middle rows of each day (location, teacher)
    DAY_END = 3  # Last row of each day (assistant)


class ColorScheme:
    """Color constants for the document"""
    DAYS_COLUMN = "8DB3E2"
    CATEGORIES_COLUMN = "B7DDE8"
    HEADER_BACKGROUND = "8DB3E2"
    SEPARATOR_BACKGROUND = "B7DDE8"


class BorderWidth:
    """Border width constants in Word units (1/8 point)"""
    THIN = 4      # 0.5pt
    THICK = 18    # 2.25pt


class FontConfig:
    """Font configuration constants"""
    FONT_NAME = 'Arial'
    TITLE_SIZE = Pt(18)
    HEADER_SIZE = Pt(12)
    FOOTER_SIZE = Pt(10)
    TABLE_CELL_SIZE = Pt(8)


class TableDimensions:
    """Table dimension constants"""
    # Page dimensions
    A4_LANDSCAPE_WIDTH_INCHES = 11.69
    A4_LANDSCAPE_HEIGHT_INCHES = 8.27
    MARGIN_INCHES = 0.5
    AVAILABLE_WIDTH_INCHES = A4_LANDSCAPE_WIDTH_INCHES - (2 * MARGIN_INCHES)
    
    # Table structure
    TOTAL_ROWS = 21  # 5 days * 4 categories + 1 header
    TOTAL_COLUMNS = 13  # 6 data columns + 3 separators + 4 half slots
    
    # Column widths (in inches)
    DAYS_COLUMN_WIDTH = 0.8
    CATEGORIES_COLUMN_WIDTH = 1.0
    TIME_SLOT_WIDTH = 1.8
    TIME_SLOT_HALF_WIDTH = 0.2  # Half width for half slots
    SEPARATOR_WIDTH = 0.2
    
    # Row structure
    HEADER_ROW_INDEX = 0
    CONTENT_START_ROW_INDEX = 1
    ROWS_PER_DAY = 4  # 4 categories per day


class WordGenerator:
    """Generates Word document with schedule table"""
    
    def __init__(self):
        self.day_mapping = {
            "الأحد": "sunday",
            "الاثنين": "monday", 
            "الثلاثاء": "tuesday",
            "الأربعاء": "wednesday",
            "الخميس": "thursday"
        }
        self.days_arabic = ["الأحد", "الاثنين", "الثلاثاء", "الأربعاء", "الخميس"]
        self.detail_categories = [
            "اسم المادة",
            "المكان", 
            "استاذ المادة",
            "الهيئة المعاونة"
        ]
        self.time_slots = [
            "المحاضرة الأولى 8.50-10.30",
            "المحاضرة الثانية 10.40 - 12.10",
            "المحاضرة الثالثة 12.20 - 1.50",
            "المحاضرة الرابعة 2.00 - 3.30"
        ]
        
        # Define column positions for better readability
        self.time_slot_positions = [
            ColumnType.TIME_SLOT_1.value,
            ColumnType.TIME_SLOT_2.value,
            ColumnType.TIME_SLOT_3.value,
            ColumnType.TIME_SLOT_4.value
        ]
        self.time_slot_half_positions = [
            ColumnType.TIME_SLOT_1_HALF.value,
            ColumnType.TIME_SLOT_2_HALF.value,
            ColumnType.TIME_SLOT_3_HALF.value,
            ColumnType.TIME_SLOT_4_HALF.value
        ]
        self.separator_positions = [
            ColumnType.SEPARATOR_1.value,
            ColumnType.SEPARATOR_2.value,
            ColumnType.SEPARATOR_3.value
        ]
        self.slot_keys = ["first", "second", "third", "fourth"]
        
        # Column width mapping
        self.column_widths = {
            ColumnType.DAYS.value: TableDimensions.DAYS_COLUMN_WIDTH,
            ColumnType.CATEGORIES.value: TableDimensions.CATEGORIES_COLUMN_WIDTH,
            ColumnType.TIME_SLOT_1.value: TableDimensions.TIME_SLOT_WIDTH,
            ColumnType.TIME_SLOT_1_HALF.value: TableDimensions.TIME_SLOT_HALF_WIDTH,
            ColumnType.SEPARATOR_1.value: TableDimensions.SEPARATOR_WIDTH,
            ColumnType.TIME_SLOT_2.value: TableDimensions.TIME_SLOT_WIDTH,
            ColumnType.TIME_SLOT_2_HALF.value: TableDimensions.TIME_SLOT_HALF_WIDTH,
            ColumnType.SEPARATOR_2.value: TableDimensions.SEPARATOR_WIDTH,
            ColumnType.TIME_SLOT_3.value: TableDimensions.TIME_SLOT_WIDTH,
            ColumnType.TIME_SLOT_3_HALF.value: TableDimensions.TIME_SLOT_HALF_WIDTH,
            ColumnType.SEPARATOR_3.value: TableDimensions.SEPARATOR_WIDTH,
            ColumnType.TIME_SLOT_4.value: TableDimensions.TIME_SLOT_WIDTH,
            ColumnType.TIME_SLOT_4_HALF.value: TableDimensions.TIME_SLOT_HALF_WIDTH
        }
    
    def _validate_level(self, level: str) -> str:
        """Validate that the level is one of the allowed values [100, 200, 300, 400]"""
        allowed_levels = ["100", "200", "300", "400"]
        if level not in allowed_levels:
            raise ValueError(f"Level must be one of {allowed_levels}, got: {level}")
        return level
    
    def _get_level_config(self, level: str, specialty: str = None) -> dict:
        """Get the configuration for a specific level and specialty"""
        validated_level = self._validate_level(level)
        return _get_level_config(validated_level, specialty)
    
    def create_document(self) -> Document:
        """Create a new Word document"""
        doc = Document()
        
        # Set document to landscape A4
        sections = doc.sections
        for section in sections:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(TableDimensions.A4_LANDSCAPE_WIDTH_INCHES)
            section.page_height = Inches(TableDimensions.A4_LANDSCAPE_HEIGHT_INCHES)
            section.top_margin = Inches(TableDimensions.MARGIN_INCHES)
            section.bottom_margin = Inches(TableDimensions.MARGIN_INCHES)
            section.left_margin = Inches(TableDimensions.MARGIN_INCHES)
            section.right_margin = Inches(TableDimensions.MARGIN_INCHES)
        
        return doc
    
    def add_page_header(self, doc: Document, speciality: str, level: str) -> None:
        """Add page header with identity information"""
        section = doc.sections[0]
        self._add_header_to_section(section, speciality, level)
    
    def add_page_footer(self, doc: Document, level: str, specialty: str = None) -> None:
        """Add page footer with generation information"""
        section = doc.sections[0]
        self._add_footer_to_section(section, level, specialty)
    
    def add_speciality_level_title(self, doc: Document, speciality: str, level: str) -> None:
        """Add title for specialty and level combination"""
        # Add spacing before title
        doc.add_paragraph()
        
        # Create title paragraph
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add title text
        title_text = f"جدول {_get_mapped_speciality(speciality, level)} - {_get_mapped_level(level)}"
        title_run = title_para.add_run(title_text)
        title_run.font.name = FontConfig.FONT_NAME
        title_run.font.size = FontConfig.TITLE_SIZE
        title_run.font.bold = True
        
        # Set title to RTL
        self._set_paragraph_rtl(title_para)
        
        # Add spacing after title
        doc.add_paragraph()
    
    def create_table_structure(self, doc: Document, weekly_schedule: WeeklySchedule) -> None:
        """Create the main table structure"""
        # Create table with defined dimensions
        table = doc.add_table(rows=TableDimensions.TOTAL_ROWS, cols=TableDimensions.TOTAL_COLUMNS)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set table width and disable autofit
        table.autofit = False
        table.allow_autofit = False
        
        # Calculate total width from column definitions
        total_width_inches = sum(self.column_widths.values())
        table.width = Inches(total_width_inches)
        
        # Set individual column widths using XML properties
        self._set_table_column_widths(table)
        
        # Set table to RTL
        self._set_table_rtl(table)
        
        self._fill_header_row(table)
        self._fill_content_rows(table, weekly_schedule)
        self._apply_formatting(table)
        self._apply_table_outline_borders(table)
    
    def _fill_header_row(self, table) -> None:
        """Fill the header row with time slots"""
        header_row = table.rows[TableDimensions.HEADER_ROW_INDEX]
        
        # First two cells are empty in header
        header_row.cells[ColumnType.DAYS.value].text = ""
        header_row.cells[ColumnType.CATEGORIES.value].text = ""
        
        # Fill each time slot in separate columns with separators
        for i, time_slot in enumerate(self.time_slots):
            # Main time slot column
            main_cell = header_row.cells[self.time_slot_positions[i]]
            half_cell = header_row.cells[self.time_slot_half_positions[i]]
            
            # Set content in main cell and merge with half cell
            main_cell.text = time_slot
            main_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_rtl(main_cell)
            
            # Merge main and half slot columns for header
            main_cell.merge(half_cell)
            # Clean up any extra paragraphs that might cause newlines
            if len(main_cell.paragraphs) > 1:
                # Keep only the first paragraph
                for i in range(len(main_cell.paragraphs) - 1, 0, -1):
                    main_cell.paragraphs[i]._element.getparent().remove(main_cell.paragraphs[i]._element)
            # Re-apply formatting after cleanup
            main_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_rtl(main_cell)
        
        # Set separator columns to empty
        for pos in self.separator_positions:
            header_row.cells[pos].text = ""
    
    def _fill_content_rows(self, table, weekly_schedule: WeeklySchedule) -> None:
        """Fill the content rows with schedule data"""
        row_index = TableDimensions.CONTENT_START_ROW_INDEX
        
        for day_index, day_arabic in enumerate(self.days_arabic):
            day_key = list(weekly_schedule.keys())[day_index]
            day_schedule = weekly_schedule[day_key]
            
            # For each day, create 4 rows (one for each detail category)
            for category_index, category in enumerate(self.detail_categories):
                row = table.rows[row_index]
                
                # Day column (merged vertically for each day)
                if category_index == 0:
                    day_cell = row.cells[ColumnType.DAYS.value]
                    day_cell.text = day_arabic
                    day_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._set_cell_rtl(day_cell)
                    # Merge vertically with next 3 rows
                    if row_index + 3 < len(table.rows):
                        day_cell.merge(table.rows[row_index + 3].cells[ColumnType.DAYS.value])
                        # Reapply thick borders to merged day cell
                        self._apply_day_cell_borders(day_cell, row_index)
                
                # Category column
                category_cell = row.cells[ColumnType.CATEGORIES.value]
                category_cell.text = category
                category_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._set_cell_rtl(category_cell)
                
                # Time slot columns with separators
                for slot_index, slot_key in enumerate(self.slot_keys):
                    main_cell = row.cells[self.time_slot_positions[slot_index]]
                    half_cell = row.cells[self.time_slot_half_positions[slot_index]]
                    schedule_entry = day_schedule[slot_key]
                    
                    if schedule_entry:
                        # Determine content based on category
                        if category_index == 0:  # Course name
                            content = schedule_entry.course_name
                        elif category_index == 1:  # Location
                            content = schedule_entry.location
                        elif category_index == 2:  # Instructor
                            content = schedule_entry.instructor
                        elif category_index == 3:  # Assistant
                            content = schedule_entry.assistant
                        else:
                            content = ""
                        
                        # Handle half slot logic
                        if schedule_entry.is_half_slot:
                            # For half slots: put content in main column, leave half column empty
                            main_cell.text = content
                            half_cell.text = ""
                        else:
                            # For full slots: merge the cells and put content across both columns
                            main_cell.text = content
                            half_cell.text = ""
                            # Merge the cells horizontally (main + half slot columns)
                            main_cell.merge(half_cell)
                            # Clean up any extra paragraphs that might cause newlines
                            if len(main_cell.paragraphs) > 1:
                                # Keep only the first paragraph
                                for i in range(len(main_cell.paragraphs) - 1, 0, -1):
                                    main_cell.paragraphs[i]._element.getparent().remove(main_cell.paragraphs[i]._element)
                            # Re-apply formatting after cleanup
                            main_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            self._set_cell_rtl(main_cell)
                    else:
                        # Empty slot: merge the cells and leave both empty
                        main_cell.text = ""
                        half_cell.text = ""
                        # Merge the cells horizontally (main + half slot columns)
                        main_cell.merge(half_cell)
                        # Clean up any extra paragraphs that might cause newlines
                        if len(main_cell.paragraphs) > 1:
                            # Keep only the first paragraph
                            for i in range(len(main_cell.paragraphs) - 1, 0, -1):
                                main_cell.paragraphs[i]._element.getparent().remove(main_cell.paragraphs[i]._element)
                        # Re-apply formatting after cleanup
                        main_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        self._set_cell_rtl(main_cell)
                
                # Set separator columns to empty
                for pos in self.separator_positions:
                    row.cells[pos].text = ""
                
                row_index += 1
    
    def _apply_formatting(self, table) -> None:
        """Apply formatting to the table"""
        # Apply borders and formatting to all cells
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                # Set font
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = FontConfig.FONT_NAME
                        run.font.size = FontConfig.TABLE_CELL_SIZE
                
                # Set cell borders with different widths based on row and column types
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Determine row type
                if row_index == TableDimensions.HEADER_ROW_INDEX:
                    row_type = RowType.HEADER
                else:
                    # Calculate which row within the day this is (0-3)
                    day_row_index = (row_index - TableDimensions.CONTENT_START_ROW_INDEX) % TableDimensions.ROWS_PER_DAY
                    if day_row_index == 0:
                        row_type = RowType.DAY_START
                    elif day_row_index == TableDimensions.ROWS_PER_DAY - 1:
                        row_type = RowType.DAY_END
                    else:
                        row_type = RowType.DAY_MIDDLE
                
                # Ensure thick vertical borders for the day names and categories column (outer-left outline)
                vertical_border_width = str(BorderWidth.THICK)
                
                # Horizontal borders depend on row type
                if row_type == RowType.HEADER:
                    # Header row – make top border thick to complete outer outline
                    top_border_width = str(BorderWidth.THICK)
                    bottom_border_width = str(BorderWidth.THIN)
                elif row_type == RowType.DAY_START:
                    # First row of day - top border thick, bottom border thin
                    top_border_width = str(BorderWidth.THICK)
                    # If this is the merged day cell row, we also want a thick bottom
                    # border to ensure the outline of the merged cell across rows.
                    if col_index == ColumnType.DAYS.value:
                        bottom_border_width = str(BorderWidth.THICK)
                    else:
                        bottom_border_width = str(BorderWidth.THIN)
                elif row_type == RowType.DAY_END:
                    # Last row of day - top border thin, bottom border thick
                    top_border_width = str(BorderWidth.THIN)
                    bottom_border_width = str(BorderWidth.THICK)
                else:  # DAY_MIDDLE
                    # Middle rows of day - all borders thin
                    top_border_width = str(BorderWidth.THIN)
                    bottom_border_width = str(BorderWidth.THIN)
                
                tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                                    f'<w:top w:val="single" w:sz="{top_border_width}" w:space="0" w:color="000000"/>'
                                    f'<w:left w:val="single" w:sz="{vertical_border_width}" w:space="0" w:color="000000"/>'
                                    f'<w:bottom w:val="single" w:sz="{bottom_border_width}" w:space="0" w:color="000000"/>'
                                    f'<w:right w:val="single" w:sz="{vertical_border_width}" w:space="0" w:color="000000"/>'
                                    f'</w:tcBorders>')
                tcPr.append(tcBorders)
                
                # Apply background colors based on column type and row position
                self._apply_cell_background_color(tcPr, col_index, row_index)
                
                # Ensure all cells have RTL formatting and center alignment
                if cell.text.strip():  # Only apply to non-empty cells
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._set_cell_rtl(cell)
    
    def _apply_day_cell_borders(self, day_cell, row_index: int) -> None:
        """Apply thick borders to merged day cells after merging"""
        tc = day_cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Remove any existing borders
        for child in tcPr:
            if child.tag.endswith('tcBorders'):
                tcPr.remove(child)
        
        # For merged cells, we need to ensure the borders apply to the entire merged region
        # Check if this is a merged cell by looking for vMerge element
        is_merged = False
        for child in tcPr:
            if child.tag.endswith('vMerge'):
                is_merged = True
                break
        
        # Apply thick borders all around the merged day cell
        # For merged cells, we need to be more explicit about the border application
        if is_merged:
            # Use more explicit border settings for merged cells
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                                f'<w:top w:val="single" w:sz="{BorderWidth.THICK}" w:space="0" w:color="000000"/>'
                                f'<w:left w:val="single" w:sz="{BorderWidth.THICK}" w:space="0" w:color="000000"/>'
                                f'<w:bottom w:val="single" w:sz="{BorderWidth.THICK}" w:space="0" w:color="000000"/>'
                                f'<w:right w:val="single" w:sz="{BorderWidth.THICK}" w:space="0" w:color="000000"/>'
                                f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                                f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                                f'</w:tcBorders>')
        else:
            # Standard border application for non-merged cells
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                                f'<w:top w:val="single" w:sz="{BorderWidth.THICK}" w:space="0" w:color="000000"/>'
                                f'<w:left w:val="single" w:sz="{BorderWidth.THICK}" w:space="0" w:color="000000"/>'
                                f'<w:bottom w:val="single" w:sz="{BorderWidth.THICK}" w:space="0" w:color="000000"/>'
                                f'<w:right w:val="single" w:sz="{BorderWidth.THICK}" w:space="0" w:color="000000"/>'
                                f'</w:tcBorders>')
        
        tcPr.append(tcBorders)
        
        # Reapply background color for day column
        self._apply_background_color(tcPr, ColorScheme.DAYS_COLUMN)
        
        # Ensure font formatting is maintained
        for paragraph in day_cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = FontConfig.FONT_NAME
                run.font.size = FontConfig.TABLE_CELL_SIZE
        
        # For merged cells, also ensure the table-level borders are properly set
        if is_merged:
            self._ensure_merged_cell_table_borders(day_cell)
    
    def _ensure_merged_cell_table_borders(self, merged_cell) -> None:
        """Ensure table-level borders are properly set for merged cells"""
        # Get the table that contains this cell
        table = merged_cell._tc.getparent().getparent()
        
        # Get or create table properties
        tblPr = table.tblPr
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            table.insert(0, tblPr)
        
        # Remove any existing table borders
        for child in tblPr:
            if child.tag.endswith('tblBorders'):
                tblPr.remove(child)
        
        # Add table-level borders that ensure merged cells respect the border rules
        tblBorders = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
                             f'<w:top w:val="single" w:sz="{BorderWidth.THICK}" w:space="0" w:color="000000"/>'
                             f'<w:left w:val="single" w:sz="{BorderWidth.THICK}" w:space="0" w:color="000000"/>'
                             f'<w:bottom w:val="single" w:sz="{BorderWidth.THICK}" w:space="0" w:color="000000"/>'
                             f'<w:right w:val="single" w:sz="{BorderWidth.THICK}" w:space="0" w:color="000000"/>'
                             f'<w:insideH w:val="single" w:sz="{BorderWidth.THIN}" w:space="0" w:color="000000"/>'
                             f'<w:insideV w:val="single" w:sz="{BorderWidth.THIN}" w:space="0" w:color="000000"/>'
                             f'</w:tblBorders>')
        tblPr.append(tblBorders)
    
    def _apply_cell_background_color(self, tcPr, col_index: int, row_index: int) -> None:
        """Apply background color to a table cell based on its position"""
        # Day names column (except header row)
        if col_index == ColumnType.DAYS.value and row_index > TableDimensions.HEADER_ROW_INDEX:
            self._apply_background_color(tcPr, ColorScheme.DAYS_COLUMN)
        
        # Detail categories column (except header row)
        elif col_index == ColumnType.CATEGORIES.value and row_index > TableDimensions.HEADER_ROW_INDEX:
            self._apply_background_color(tcPr, ColorScheme.CATEGORIES_COLUMN)
        
        # Time slots in header row (both main and half slot columns)
        elif row_index == TableDimensions.HEADER_ROW_INDEX and (col_index in self.time_slot_positions or col_index in self.time_slot_half_positions):
            self._apply_background_color(tcPr, ColorScheme.HEADER_BACKGROUND)
        
        # Separator columns in header row
        elif row_index == TableDimensions.HEADER_ROW_INDEX and col_index in self.separator_positions:
            self._apply_background_color(tcPr, ColorScheme.HEADER_BACKGROUND)
        
        # Separator columns in content rows
        elif row_index > TableDimensions.HEADER_ROW_INDEX and col_index in self.separator_positions:
            self._apply_background_color(tcPr, ColorScheme.SEPARATOR_BACKGROUND)
    
    def _apply_table_outline_borders(self, table) -> None:
        """Apply thick borders to the entire table outline"""
        rows = table.rows
        total_rows = len(rows)
        total_cols = len(rows[0].cells) if rows else 0
        
        # Get the Word namespace URI
        w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        
        for row_index, row in enumerate(rows):
            for col_index, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Check if this cell is on the table outline
                is_top_edge = row_index == 0
                is_bottom_edge = row_index == total_rows - 1
                is_left_edge = col_index == 0
                is_right_edge = col_index == total_cols - 1
                
                # If cell is on any edge, ensure that edge has thick border
                if is_top_edge or is_bottom_edge or is_left_edge or is_right_edge:
                    # Get existing borders or create new ones
                    existing_borders = None
                    for child in tcPr:
                        if child.tag.endswith('tcBorders'):
                            existing_borders = child
                            break
                    
                    if existing_borders is not None:
                        # Update existing borders for outline edges
                        if is_top_edge:
                            for border in existing_borders:
                                if border.tag.endswith('top'):
                                    border.set(f'{w_ns}sz', str(BorderWidth.THICK))
                        if is_bottom_edge:
                            for border in existing_borders:
                                if border.tag.endswith('bottom'):
                                    border.set(f'{w_ns}sz', str(BorderWidth.THICK))
                        if is_left_edge:
                            for border in existing_borders:
                                if border.tag.endswith('left'):
                                    border.set(f'{w_ns}sz', str(BorderWidth.THICK))
                        if is_right_edge:
                            for border in existing_borders:
                                if border.tag.endswith('right'):
                                    border.set(f'{w_ns}sz', str(BorderWidth.THICK))
    
    def _apply_background_color(self, tcPr, color_hex: str) -> None:
        """Apply background color to a table cell"""
        # Remove any existing shading
        for child in tcPr:
            if child.tag.endswith('shd'):
                tcPr.remove(child)
        
        # Add shading with the specified color
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        tcPr.append(shd)
    
    def _set_table_column_widths(self, table) -> None:
        """Set column widths for a table"""
        for i, column in enumerate(table.columns):
            # Get width for this column
            width = Inches(self.column_widths[i])
            
            # Apply width to all cells in the column
            for cell in column.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Remove any existing width setting
                for child in tcPr:
                    if child.tag.endswith('tcW'):
                        tcPr.remove(child)
                
                # Add width property
                tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(width.inches * 1440)}" w:type="dxa"/>')
                tcPr.append(tcW)
    
    def _set_table_rtl(self, table) -> None:
        """Set the entire table to RTL direction"""
        tbl = table._tbl
        
        # Get or create tblPr element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tbl.insert(0, tblPr)
        
        # Remove any existing bidiVisual setting
        for child in tblPr:
            if child.tag.endswith('bidiVisual'):
                tblPr.remove(child)
        
        # Add RTL setting
        bidiVisual = parse_xml(f'<w:bidiVisual {nsdecls("w")}/>')
        tblPr.append(bidiVisual)
    
    def _set_cell_rtl(self, cell) -> None:
        """Set a cell to RTL direction"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Remove any existing textDirection setting
        for child in tcPr:
            if child.tag.endswith('textDirection'):
                tcPr.remove(child)
        
        # Add RTL text direction (horizontal RTL)
        textDirection = parse_xml(f'<w:textDirection {nsdecls("w")} w:val="rtl"/>')
        tcPr.append(textDirection)
        
        # Set paragraph direction to RTL
        for paragraph in cell.paragraphs:
            self._set_paragraph_rtl(paragraph)
    
    def _set_paragraph_rtl(self, paragraph) -> None:
        """Set a paragraph to RTL direction"""
        p = paragraph._p
        
        # Get or create pPr element
        pPr = p.pPr
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
            p.insert(0, pPr)
        
        # Remove any existing bidi setting
        for child in pPr:
            if child.tag.endswith('bidi'):
                pPr.remove(child)
        
        # Add RTL paragraph direction
        bidi = parse_xml(f'<w:bidi {nsdecls("w")}/>')
        pPr.append(bidi)
    
    def _clear_section_content(self, section_element) -> None:
        """Safely clear content from a section element (header/footer)"""
        try:
            for element in list(section_element._element):
                section_element._element.remove(element)
        except:
            # If clearing fails, continue silently
            pass
    
    def _create_header_table(self, header) -> Any:
        """Create and configure header table with proper column widths"""
        # Add header table with three columns: right, center, left
        header_table = header.add_table(rows=3, cols=3, width=Inches(TableDimensions.AVAILABLE_WIDTH_INCHES))
        header_table.style = 'Table Grid'
        header_table.autofit = False
        header_table.allow_autofit = False
        
        # Set column widths (right, center, left)
        column_widths = [3.5, 3.69, 3.5]  # in inches
        for i, column in enumerate(header_table.columns):
            for cell in column.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(column_widths[i] * 1440)}" w:type="dxa"/>')
                tcPr.append(tcW)
        
        return header_table
    
    def _fill_header_content(self, header_table, speciality: str, level: str) -> None:
        """Fill header table with content"""
        # Get level-specific configuration
        level_config = self._get_level_config(level, speciality)
        header_config = level_config["header"]
        
        # Right column (University/Faculty/Department info)
        header_table.rows[0].cells[0].text = header_config["university_name"]
        header_table.rows[1].cells[0].text = header_config["faculty_name"]
        header_table.rows[2].cells[0].text = f"{header_config['department_prefix']}"
        
        # Center column (Logo and title)
        # Add logo to the first row
        logo_cell = header_table.rows[0].cells[1]
        self._add_logo_to_cell(logo_cell)
        
        # Add title to the third row (different from original logic)
        title_cell = header_table.rows[2].cells[1]
        if title_cell.paragraphs[0].runs:
            # If logo was added, add title to a new paragraph
            title_para = title_cell.add_paragraph()
            title_para.text = header_config["schedule_template"]
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_paragraph_rtl(title_para)
        else:
            # If no logo, add title to first paragraph
            title_cell.text = header_config["schedule_template"]
            title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_rtl(title_cell)
        
        # Empty second row in center column
        header_table.rows[1].cells[1].text = ""
        
        # Left column (Academic details)
        header_table.rows[0].cells[2].text = header_config["academic_year"]
        header_table.rows[1].cells[2].text = header_config["semester"]
        header_table.rows[2].cells[2].text = f"{header_config['level_prefix']} {_get_mapped_level(level)} {header_config['division_prefix']} {_get_mapped_speciality(speciality, level)}"
    
    def _apply_header_formatting(self, header_table) -> None:
        """Apply formatting to header table"""
        for row in header_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = FontConfig.FONT_NAME
                        run.font.size = FontConfig.HEADER_SIZE
                        run.font.bold = True
                        self._set_paragraph_rtl(paragraph)
        
        # Set header table to RTL
        self._set_table_rtl(header_table)
        
        # Remove borders from header table to match image
        self._remove_table_borders(header_table)
    
    def _add_header_to_section(self, section, speciality: str, level: str) -> None:
        """Add header to a specific section"""
        # Create header
        header = section.header
        
        # Clear existing content safely
        self._clear_section_content(header)
        
        # Create and configure header table
        header_table = self._create_header_table(header)
        
        # Fill header content
        self._fill_header_content(header_table, speciality, level)
        
        # Apply formatting
        self._apply_header_formatting(header_table)
    
    def _add_footer_to_section(self, section, level: str, specialty: str = None) -> None:
        """Add footer to a specific section"""
        # Create footer
        footer = section.footer
        
        # Clear existing content safely
        self._clear_section_content(footer)
        
        # Create footer table with dynamic columns based on level
        footer_table = self._create_footer_table(footer, level, specialty)
        
        # Fill footer content
        self._fill_footer_content(footer_table, level, specialty)
        
        # Apply formatting
        self._apply_footer_formatting(footer_table)
        
        # Add generation info below the table
        self._add_generation_info(footer)
    
    def _add_logo_to_cell(self, cell) -> None:
        """Add logo image to a table cell"""
        try:
            # Get the path to the logo file
            current_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            logo_path = os.path.join(current_dir, "assets", "logo.png")
            
            # Check if logo file exists
            if os.path.exists(logo_path):
                # Clear any existing content in the cell
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = ""
                
                # Add the logo image to the first paragraph
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add the image with appropriate size (smaller to match image)
                run = paragraph.add_run()
                run.add_picture(logo_path, width=Inches(0.98), height=Inches(0.55))
                
                # Set RTL for the paragraph
                self._set_paragraph_rtl(paragraph)
            else:
                # If logo doesn't exist, add placeholder text
                cell.text = "[LOGO]"
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._set_cell_rtl(cell)
                
        except Exception as e:
            # If there's any error, add placeholder text
            cell.text = "[LOGO]"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_rtl(cell)
    
    def _create_footer_table(self, footer, level: str, specialty: str = None) -> Any:
        """Create and configure footer table with proper column widths"""
        level_config = self._get_level_config(level, specialty)
        footer_config = level_config["footer"]
        
        # Check if program manager is present (levels 100, 200 have program manager)
        has_program_manager = "program_manager_title" in footer_config
        
        if has_program_manager:
            cols = 4
            column_widths = [2.67, 2.67, 2.67, 2.68]  # Equal width for 4 columns
        else:
            cols = 3
            column_widths = [3.5, 3.69, 3.5]  # in inches
        
        # Add footer table with dynamic columns
        footer_table = footer.add_table(rows=2, cols=cols, width=Inches(TableDimensions.AVAILABLE_WIDTH_INCHES))
        footer_table.style = 'Table Grid'
        footer_table.autofit = False
        footer_table.allow_autofit = False
        
        # Set column widths
        for i, column in enumerate(footer_table.columns):
            for cell in column.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(column_widths[i] * 1440)}" w:type="dxa"/>')
                tcPr.append(tcW)
        
        return footer_table
    
    def _fill_footer_content(self, footer_table, level: str, specialty: str = None) -> None:
        """Fill footer table with content"""
        # Get level-specific configuration
        level_config = self._get_level_config(level, specialty)
        footer_config = level_config["footer"]
        
        # Check if program manager is present (levels 100, 200 have program manager)
        has_program_manager = "program_manager_title" in footer_config
        
        if has_program_manager:
            # For levels 100, 200, use 4 columns including program manager
            # Column 0: Dean info
            footer_table.rows[0].cells[0].text = footer_config['dean_title']
            footer_table.rows[1].cells[0].text = footer_config['dean_name']
            
            # Column 1: Vice Dean info
            footer_table.rows[0].cells[1].text = footer_config['vice_dean_title']
            footer_table.rows[1].cells[1].text = footer_config['vice_dean_name']
            
            # Column 2: Head of Department info
            footer_table.rows[0].cells[2].text = footer_config['head_department_title']
            footer_table.rows[1].cells[2].text = footer_config['head_department_name']
            
            # Column 3: Program Manager info
            footer_table.rows[0].cells[3].text = footer_config['program_manager_title']
            footer_table.rows[1].cells[3].text = footer_config['program_manager_name']
        else:
            # For levels 300, 400, use 3 columns
            # Right column (Dean info)
            footer_table.rows[0].cells[0].text = footer_config['dean_title']
            footer_table.rows[1].cells[0].text = footer_config['dean_name']
            
            # Center column (Vice Dean info)
            footer_table.rows[0].cells[1].text = footer_config['vice_dean_title']
            footer_table.rows[1].cells[1].text = footer_config['vice_dean_name']
            
            # Left column (Head of Department info)
            footer_table.rows[0].cells[2].text = footer_config['head_department_title']
            footer_table.rows[1].cells[2].text = footer_config['head_department_name']
    
    def _apply_footer_formatting(self, footer_table) -> None:
        """Apply formatting to footer table"""
        for row in footer_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = FontConfig.FONT_NAME
                        run.font.size = FontConfig.FOOTER_SIZE
                        run.font.bold = True
                        self._set_paragraph_rtl(paragraph)
        
        # Set footer table to RTL
        self._set_table_rtl(footer_table)
        
        # Remove borders from footer table to match image
        self._remove_table_borders(footer_table)
    
    def _add_generation_info(self, footer) -> None:
        """Add generation information below the footer table"""
        # Add spacing
        
        
        # Add generation info paragraph
        footer_para = footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add footer text
        current_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        footer_text = FOOTER_GENERATION_INFO.format(date=current_date)
        footer_run = footer_para.add_run(footer_text)
        footer_run.font.name = FontConfig.FONT_NAME
        footer_run.font.size = FontConfig.FOOTER_SIZE
        footer_run.font.italic = True
        
        # Set footer to RTL
        self._set_paragraph_rtl(footer_para)
    
    def _remove_table_borders(self, table) -> None:
        """Remove all borders from a table"""
        tbl = table._tbl
        
        # Get or create tblPr element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tbl.insert(0, tblPr)
        
        # Remove any existing tblBorders setting
        for child in tblPr:
            if child.tag.endswith('tblBorders'):
                tblPr.remove(child)
        
        # Add no borders setting
        tblBorders = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
                             f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                             f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                             f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                             f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                             f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                             f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                             f'</w:tblBorders>')
        tblPr.append(tblBorders)
    
    def generate_word_document(self, weekly_schedule: WeeklySchedule, output_path: str) -> None:
        """Generate Word document from weekly schedule (backward compatibility)"""
        doc = self.create_document()
        self.create_table_structure(doc, weekly_schedule)
        doc.save(output_path)
    
    def generate_multi_level_word_document(self, multi_level_schedule: MultiLevelSchedule, output_path: str) -> None:
        """Generate Word document with multiple tables for each specialty-level combination"""
        doc = self.create_document()
        
        # Generate table for each specialty-level combination with sections
        for i, schedule in enumerate(multi_level_schedule.schedules):
            print(f"📄 Generating table for {schedule.speciality} - {schedule.level}")
            
            # Create new section for each page (except the first one)
            if i > 0:
                # Add page break
                doc.add_page_break()
                # Create new section
                new_section = doc.add_section()
                # Break the link to previous section's header
                new_section.header.is_linked_to_previous = False
                new_section.footer.is_linked_to_previous = False
                # Add header and footer to the new section
                self._add_header_to_section(new_section, schedule.speciality, schedule.level)
                self._add_footer_to_section(new_section, schedule.level, schedule.speciality)
            else:
                # First page uses the default section
                section = doc.sections[0]
                # Break the link to previous section's header (for first section, this ensures it's independent)
                section.header.is_linked_to_previous = False
                section.footer.is_linked_to_previous = False
                self._add_header_to_section(section, schedule.speciality, schedule.level)
                self._add_footer_to_section(section, schedule.level, schedule.speciality)
            
            # Create table for this combination
            doc.add_paragraph()
            self.create_table_structure(doc, schedule.weekly_schedule)
        
        doc.save(output_path)


class LocationWordGenerator(WordGenerator):
    """Generates Word document with schedule table for each location"""

    def _fill_header_content(self, header_table, location: str) -> None:
        """Fill header table with content for location view"""
        # Get base configuration
        header_config = BASE_LEVEL_CONFIG["header"]

        # Right column (University/Faculty/Department info)
        header_table.rows[0].cells[0].text = header_config["university_name"]
        header_table.rows[1].cells[0].text = header_config["faculty_name"]
        header_table.rows[2].cells[0].text = f"{header_config['department_prefix']}"

        # Center column (Logo and title)
        logo_cell = header_table.rows[0].cells[1]
        self._add_logo_to_cell(logo_cell)

        title_cell = header_table.rows[2].cells[1]
        title_cell.text = header_config["schedule_template"]
        title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_cell_rtl(title_cell)

        header_table.rows[1].cells[1].text = ""

        # Left column (Academic details)
        header_table.rows[0].cells[2].text = header_config["academic_year"]
        header_table.rows[1].cells[2].text = header_config["semester"]
        header_table.rows[2].cells[2].text = f"جدول إشغال {location}"

    def _add_header_to_section(self, section, location: str) -> None:
        """Add header to a specific section for location view"""
        header = section.header
        self._clear_section_content(header)
        header_table = self._create_header_table(header)
        self._fill_header_content(header_table, location)
        self._apply_header_formatting(header_table)

    def add_location_title(self, doc: Document, location: str) -> None:
        """Add title for location"""
        doc.add_paragraph()
        return
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_text = f"جدول إشغال {location}"
        title_run = title_para.add_run(title_text)
        title_run.font.name = FontConfig.FONT_NAME
        title_run.font.size = FontConfig.TITLE_SIZE
        title_run.font.bold = True
        self._set_paragraph_rtl(title_para)
        doc.add_paragraph()

    def generate_multi_location_word_document(self, multi_location_schedule: MultiLocationSchedule, output_path: str) -> None:
        """Generate Word document with multiple tables for each location"""
        doc = self.create_document()

        for i, schedule in enumerate(multi_location_schedule.schedules):
            print(f"📄 Generating table for {schedule.location}")

            if i > 0:
                doc.add_page_break()
                new_section = doc.add_section()
                new_section.header.is_linked_to_previous = False
                new_section.footer.is_linked_to_previous = False
                self._add_header_to_section(new_section, schedule.location)
                # Using a generic footer, since level and specialty are not available
                self._add_footer_to_section(new_section, "100", "comm")
            else:
                section = doc.sections[0]
                section.header.is_linked_to_previous = False
                section.footer.is_linked_to_previous = False
                self._add_header_to_section(section, schedule.location)
                # Using a generic footer, since level and specialty are not available
                self._add_footer_to_section(section, "100", "comm")

            self.add_location_title(doc, schedule.location)
            self.create_table_structure(doc, schedule.weekly_schedule)

        doc.save(output_path)


class StaffWordGenerator(WordGenerator):
    """Generates Word document with schedule table for each staff member"""

    def add_staff_title(self, doc: Document, staff_name: str) -> None:
        """Add title for staff member"""
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_para.paragraph_format.space_after = Pt(0)
        title_para.paragraph_format.keep_with_next = True
        title_text = f"جدول احمال  {staff_name}"
        title_run = title_para.add_run(title_text)
        title_run.font.name = FontConfig.FONT_NAME
        title_run.font.size = Pt(10)
        title_run.font.bold = True
        self._set_paragraph_rtl(title_para)

    def create_staff_table_structure(self, doc: Document, weekly_schedule: WeeklySchedule) -> None:
        """Create the main table structure for a staff member, filtering out empty days"""
        active_days = [day for day, day_schedule in weekly_schedule.items() if any(day_schedule.values())]
        if not active_days:
            return  # No table needed if no active days

        num_rows = len(active_days) * TableDimensions.ROWS_PER_DAY + 1
        table = doc.add_table(rows=num_rows, cols=TableDimensions.TOTAL_COLUMNS)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False

        total_width_inches = sum(self.column_widths.values())
        table.width = Inches(total_width_inches)
        self._set_table_column_widths(table)
        self._set_table_rtl(table)

        self._fill_header_row(table)
        self._fill_staff_content_rows(table, weekly_schedule, active_days)
        self._apply_formatting(table)
        self._apply_table_outline_borders(table)

    def _fill_staff_content_rows(self, table, weekly_schedule: WeeklySchedule, active_days: List[str]) -> None:
        """Fill the content rows with schedule data for active days"""
        row_index = TableDimensions.CONTENT_START_ROW_INDEX
        day_map = {v: k for k, v in self.day_mapping.items()}

        for day_key in active_days:
            day_arabic = day_map.get(day_key, "")
            day_schedule = weekly_schedule[day_key]

            for category_index, category in enumerate(self.detail_categories):
                row = table.rows[row_index]
                if category_index == 0:
                    day_cell = row.cells[ColumnType.DAYS.value]
                    day_cell.text = day_arabic
                    day_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._set_cell_rtl(day_cell)
                    if row_index + 3 < len(table.rows):
                        day_cell.merge(table.rows[row_index + 3].cells[ColumnType.DAYS.value])
                        self._apply_day_cell_borders(day_cell, row_index)

                category_cell = row.cells[ColumnType.CATEGORIES.value]
                category_cell.text = category
                category_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._set_cell_rtl(category_cell)

                for slot_index, slot_key in enumerate(self.slot_keys):
                    main_cell = row.cells[self.time_slot_positions[slot_index]]
                    half_cell = row.cells[self.time_slot_half_positions[slot_index]]
                    schedule_entry = day_schedule[slot_key]

                    if schedule_entry:
                        if category_index == 0: content = schedule_entry.course_name
                        elif category_index == 1: content = schedule_entry.location
                        elif category_index == 2: content = schedule_entry.instructor
                        elif category_index == 3: content = schedule_entry.assistant
                        else: content = ""

                        if schedule_entry.is_half_slot:
                            main_cell.text = content
                            half_cell.text = ""
                        else:
                            main_cell.text = content
                            half_cell.text = ""
                            main_cell.merge(half_cell)
                            if len(main_cell.paragraphs) > 1:
                                for i in range(len(main_cell.paragraphs) - 1, 0, -1):
                                    main_cell.paragraphs[i]._element.getparent().remove(main_cell.paragraphs[i]._element)
                            main_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            self._set_cell_rtl(main_cell)
                    else:
                        main_cell.text = ""
                        half_cell.text = ""
                        main_cell.merge(half_cell)
                        if len(main_cell.paragraphs) > 1:
                            for i in range(len(main_cell.paragraphs) - 1, 0, -1):
                                main_cell.paragraphs[i]._element.getparent().remove(main_cell.paragraphs[i]._element)
                        main_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        self._set_cell_rtl(main_cell)

                for pos in self.separator_positions:
                    row.cells[pos].text = ""
                
                row_index += 1

    def generate_multi_staff_word_document(self, multi_staff_schedule: MultiStaffSchedule, output_path: str, staff_type: str) -> None:
        """Generate Word document with multiple tables for each staff member"""
        doc = self.create_document()
        # Use a generic header/footer for staff view
        section = doc.sections[0]
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        self._add_header_to_section(section, "", "100")  # Generic header
        self._add_footer_to_section(section, "100", "comm")  # Generic footer

        MAX_TABLES_PER_PAGE = 2
        tables_on_page = 0

        for schedule in multi_staff_schedule.schedules:
            print(f"📄 Generating table for {schedule.staff_name}")

            active_days = [day for day, day_schedule in schedule.weekly_schedule.items() if any(day_schedule.values())]
            if not active_days:
                continue

            # Add page break if MAX_TABLES_PER_PAGE is reached and it's not the very first table
            if tables_on_page >= MAX_TABLES_PER_PAGE:
                doc.add_page_break()
                tables_on_page = 0

            self.add_staff_title(doc, schedule.staff_name)
            self.create_staff_table_structure(doc, schedule.weekly_schedule)
            
            tables_on_page += 1

        doc.save(output_path)