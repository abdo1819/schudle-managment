#!/usr/bin/env python3
"""
Simple test script to verify CSV to Word conversion functionality
"""

import csv
import tempfile
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def read_csv_simple(file_path):
    """Read CSV file and return data as list of dictionaries"""
    data = []
    with open(file_path, 'r', encoding='utf-8') as file:
        reader = csv.DictReader(file)
        for row in reader:
            data.append(row)
    return data


def create_schedule_structure(csv_data):
    """Create schedule structure from CSV data"""
    day_mapping = {
        "الأحد": "sunday",
        "الاثنين": "monday", 
        "الثلاثاء": "tuesday",
        "الأربعاء": "wednesday",
        "الخميس": "thursday"
    }
    
    slot_mapping = {
        "1": "first",
        "2": "second", 
        "3": "third",
        "4": "fourth"
    }
    
    # Initialize empty schedule
    schedule = {
        "sunday": {"first": None, "second": None, "third": None, "fourth": None},
        "monday": {"first": None, "second": None, "third": None, "fourth": None},
        "tuesday": {"first": None, "second": None, "third": None, "fourth": None},
        "wednesday": {"first": None, "second": None, "third": None, "fourth": None},
        "thursday": {"first": None, "second": None, "third": None, "fourth": None}
    }
    
    # Process each row
    for row in csv_data:
        day_key = day_mapping.get(row['day'])
        slot_key = slot_mapping.get(str(row['slot']))
        
        if day_key and slot_key:
            schedule[day_key][slot_key] = {
                'course_name': row['course_name'],
                'location': row['location'],
                'instructor': row['main_tutor'],
                'assistant': row['helping_stuff']
            }
    
    return schedule


def generate_word_document(schedule, output_path):
    """Generate Word document with schedule table"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Create table: 21 rows (5 days * 4 categories + 1 header) x 6 columns
    table = doc.add_table(rows=21, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    for i, column in enumerate(table.columns):
        if i == 0:  # Days column
            column.width = Inches(1.2)
        elif i == 1:  # Categories column
            column.width = Inches(1.5)
        else:  # Time slots columns
            column.width = Inches(2.0)
    
    # Fill header row
    header_row = table.rows[0]
    header_row.cells[0].text = ""
    header_row.cells[1].text = ""
    
    time_slots = [
        "المحاضرة الأولى 8.50-10.30",
        "المحاضرة الثانية 10.40 - 12.10",
        "المحاضرة الثالثة 12.20 - 1.50",
        "المحاضرة الرابعة 2.00 - 3.30"
    ]
    
    for i in range(2, 6):
        header_row.cells[i].text = time_slots[i-2]
        header_row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Merge the time slot cells
    header_cell = header_row.cells[2]
    header_cell.merge(header_row.cells[5])
    
    # Fill content rows
    days_arabic = ["الأحد", "الاثنين", "الثلاثاء", "الأربعاء", "الخميس"]
    detail_categories = [
        "اسم المادة",
        "المكان", 
        "استاذ المادة",
        "الهيئة المعاونة"
    ]
    
    row_index = 1
    for day_index, day_arabic in enumerate(days_arabic):
        day_key = list(schedule.keys())[day_index]
        day_schedule = schedule[day_key]
        
        # For each day, create 4 rows (one for each detail category)
        for category_index, category in enumerate(detail_categories):
            row = table.rows[row_index]
            
            # Day column (merged vertically for each day)
            if category_index == 0:
                row.cells[0].text = day_arabic
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Merge vertically with next 3 rows
                if row_index + 3 < len(table.rows):
                    row.cells[0].merge(table.rows[row_index + 3].cells[0])
            
            # Category column
            row.cells[1].text = category
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Time slot columns
            slot_keys = ["first", "second", "third", "fourth"]
            for slot_index, slot_key in enumerate(slot_keys):
                cell = row.cells[slot_index + 2]
                schedule_entry = day_schedule[slot_key]
                
                if schedule_entry:
                    if category_index == 0:  # Course name
                        cell.text = schedule_entry['course_name']
                    elif category_index == 1:  # Location
                        cell.text = schedule_entry['location']
                    elif category_index == 2:  # Instructor
                        cell.text = schedule_entry['instructor']
                    elif category_index == 3:  # Assistant
                        cell.text = schedule_entry['assistant']
                else:
                    cell.text = ""
            
            row_index += 1
    
    # Apply formatting
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
    
    doc.save(output_path)


def main():
    """Main function to test the conversion"""
    print("Testing CSV to Word conversion...")
    
    # Test with the sample CSV
    csv_path = "sample.csv"
    output_path = "test_output.docx"
    
    if not os.path.exists(csv_path):
        print(f"Error: {csv_path} not found")
        return
    
    try:
        # Read CSV
        print(f"Reading CSV file: {csv_path}")
        csv_data = read_csv_simple(csv_path)
        print(f"Found {len(csv_data)} rows")
        
        # Create schedule structure
        print("Creating schedule structure...")
        schedule = create_schedule_structure(csv_data)
        
        # Generate Word document
        print(f"Generating Word document: {output_path}")
        generate_word_document(schedule, output_path)
        
        print("✅ Conversion completed successfully!")
        print(f"📄 Output file: {output_path}")
        
        # Check file size
        if os.path.exists(output_path):
            size = os.path.getsize(output_path)
            print(f"📊 File size: {size} bytes")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
